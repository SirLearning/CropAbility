from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io


# ==========================================
# 1. Generate Word document (.docx)
# ==========================================

def set_cell_border(cell, **kwargs):
    """
    Helper function to set cell borders
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        if border_name in kwargs:
            val = kwargs[border_name]
            tag = 'w:{}'.format(border_name)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)

            element.set(qn('w:val'), val.get('val', 'single'))
            element.set(qn('w:sz'), str(val.get('sz', 4)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), val.get('color', 'auto'))


def create_birth_cert_docx(filename):
    doc = Document()

    # Page setup - Narrow margins
    section = doc.sections[0]
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    # Style setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)
    # Ensure Chinese characters are handled
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # --- Header ---
    header_table = doc.add_table(rows=1, cols=3)
    header_table.autofit = False
    header_table.columns[0].width = Inches(2.5)
    header_table.columns[1].width = Inches(3)  # Center title
    header_table.columns[2].width = Inches(2)

    # Left: State info
    c0 = header_table.cell(0, 0)
    p = c0.paragraphs[0]
    p.add_run("加利福尼亚州\n人口记录证明\n橙县\n卫生部门").bold = True

    # Center: Title
    c1 = header_table.cell(0, 1)
    p = c1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("加利福尼亚州出生证明")
    run.bold = True
    run.font.size = Pt(16)

    # Right: Numbers
    c2 = header_table.cell(0, 2)
    p = c2.paragraphs[0]
    p.add_run("州档案编号:\n").bold = True
    p.add_run("1052019619302\n")
    p.add_run("当地注册号:\n").bold = True
    p.add_run("1201930509213")

    doc.add_paragraph().add_run().font.size = Pt(4)  # Spacer

    # --- Main Table ---
    # We will simulate the grid.
    # The CA birth cert has specific sections. We'll use a single table and merge cells.
    # Estimated columns: 12 to allow flexible merging
    table = doc.add_table(rows=18, cols=12)
    table.style = 'Table Grid'
    table.autofit = False

    # Set fixed column widths (approximate) to avoid squishing
    col_width = Inches(7.5 / 12)
    for col in table.columns:
        col.width = col_width

    # Helper to fill data
    def fill_cell(row, col_start, col_end, label, content, align=WD_ALIGN_PARAGRAPH.LEFT):
        c = table.cell(row, col_start)
        if col_end > col_start:
            c.merge(table.cell(row, col_end))
        p = c.paragraphs[0]
        p.alignment = align
        if label:
            r1 = p.add_run(label + " ")
            r1.font.size = Pt(7)
            r1.bold = True
        if content:
            r2 = p.add_run(content)
            r2.font.size = Pt(10)
            r2.bold = True  # Content usually typed/bold
        return c

    # --- ROW 0: Warning/Headers ---
    # Usually "CERTIFICATE OF LIVE BIRTH" strip, here implied
    fill_cell(0, 0, 11, "婴儿信息", "", WD_ALIGN_PARAGRAPH.CENTER)

    # --- ROW 1: Name ---
    fill_cell(1, 0, 3, "1A. 姓名(名)", "又熙(SARAH LU)")
    fill_cell(1, 4, 7, "1B. 中间名", "-")
    fill_cell(1, 8, 11, "1C. 姓氏",
              "-")  # Document shows full name in one block or split fields; sample first name Youxi (SARAH LU).
    # Adjust based on input text near fields 1A-1C.
    # Sample names: Youxi (SARAH LU), Fei (FEI), Mei (MEI), Mei (MEI), (LU).
    # Correction based on extraction:
    # 1A (Name): Youxi (SARAH LU)
    # 1C (Last Name): implied or blank in snippet, but usually LU. Let's assume standard format.

    # --- ROW 2: Gender, Birth Type, Date ---
    fill_cell(2, 0, 2, "2. 性别", "女")
    fill_cell(2, 3, 5, "3A. 单胞胎/多胞胎", "单胞胎")
    fill_cell(2, 6, 7, "3B. 顺序", "-")
    fill_cell(2, 8, 9, "4A. 出生日期", "2019年06月03日")
    fill_cell(2, 10, 11, "4B. 出生时间", "14:03")

    # --- ROW 3: Place of Birth ---
    fill_cell(3, 0, 4, "5A. 出生地(医院名称)", "南岸全球医疗中心")
    fill_cell(3, 5, 8, "5B. 街道地址", "布里斯托南大街2701号")
    fill_cell(3, 9, 11, "5C. 城市", "圣安娜市")

    # --- ROW 4: County ---
    fill_cell(4, 0, 11, "5D. 郡/县", "橙县")

    # --- ROW 5: Parents Header ---
    fill_cell(5, 0, 11, "父母信息", "", WD_ALIGN_PARAGRAPH.CENTER)

    # --- ROW 6: Father Info ---
    fill_cell(6, 0, 3, "6A. 父亲名字", "非(FEI)")
    fill_cell(6, 4, 5, "6B. 中间名", "")
    fill_cell(6, 6, 8, "6C. 姓氏", "(LU)")  # Based on source line 46
    fill_cell(6, 9, 10, "7. 出生地", "中国")
    fill_cell(6, 11, 11, "8. 出生日期", "1980年12月18日")

    # --- ROW 7: Mother Info ---
    fill_cell(7, 0, 3, "9A. 母亲名字", "媚(MEI)")  # Source line 58
    fill_cell(7, 4, 5, "9B. 中间名", "")
    fill_cell(7, 6, 8, "9C. 姓氏", "梅(MEI)")  # Source line 59
    fill_cell(7, 9, 10, "10. 出生地", "中国")
    fill_cell(7, 11, 11, "11. 出生日期", "1985年03月18日")

    # --- ROW 8: Certification Header ---
    fill_cell(8, 0, 11, "认证", "", WD_ALIGN_PARAGRAPH.CENTER)

    # --- ROW 9: Parent Signature ---
    fill_cell(9, 0, 6, "12A. 父母/申报人签名", "非(Fei Lu) (签名)")
    fill_cell(9, 7, 9, "12B. 关系", "父亲")
    fill_cell(9, 10, 11, "12C. 签字日期", "2019年06月04日")

    # --- ROW 10: Attendant Signature ---
    fill_cell(10, 0, 6, "13A. 证明人签字", "MONICA DIAZ, ROI协调员 (签名)")
    fill_cell(10, 7, 9, "13B. 执照号码", "A45852")
    fill_cell(10, 10, 11, "13C. 签字日期", "2019年06月04日")

    # --- ROW 11: Attendant Address ---
    fill_cell(11, 0, 11, "13D. 护理人员姓名及地址", "HIEP TRUONG, 医学博士, 加州塔斯廷, 蜜蜂路16619号, 邮编:92782")

    # --- ROW 12: Other Attendant ---
    fill_cell(12, 0, 11, "14. 其他证明人", "MONICA DIAZ, ROI协调员")

    # --- ROW 13: Death Info (Usually blank) ---
    fill_cell(13, 0, 5, "15A. 死亡日期", "-")
    fill_cell(13, 6, 11, "15B. 本州存档编号", "-")

    # --- ROW 14: Local Registrar ---
    fill_cell(14, 0, 8, "16. 当地注册员签字", "NICHOLE QUICK, MD (签名)")
    fill_cell(14, 9, 11, "17. 登记日期", "2019年06月05日")

    # --- ROW 15: Footer/Certification ---
    footer_cell = fill_cell(15, 0, 11, "", "本副本真实准确，特此证明。\n签发日期: 2019年6月14日",
                            WD_ALIGN_PARAGRAPH.CENTER)

    # --- ROW 16: Seals text representation (clean) ---
    txt = "ERIC G. HANDLER, M.D.\n加州橙县卫生官员\n(此处原为卫生部印章，已省略)"
    fill_cell(16, 0, 11, "", txt, WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(filename)


# ==========================================
# 2. Generate PDF document (.pdf)
# ==========================================
# Due to font limitations in the Python environment, PDF output may not render
# Chinese characters correctly (shown as boxes), so Word is the primary output.
# For completeness, a basic PDF could be attempted, or users can export Word to PDF.
# A txt hint file could note that DOCX is the main deliverable.
# ReportLab built-in fonts can produce an English PDF if needed; Chinese requires font support.
# Given environment uncertainty, focus on DOCX — fully editable and widely supported.

create_birth_cert_docx("出生证明_无章重制版.docx")